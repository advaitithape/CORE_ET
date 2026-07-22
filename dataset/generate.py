"""Emit filled demo documents (xlsx + docx) for parts 2 & 3 from canonical data.

Produces professional-looking files under data/generated/<part_no>/ that mirror the
real DAVI / Shivkrupa QMS templates — used for the Documents view, the upload demo,
and repository credibility. The running system reads the canonical data in parts.py
directly, so these files never need to round-trip through a parser.
"""
import os
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from .parts import PARTS, INWARD, INPROCESS, PDIR, NCRS, CAPAS

OUT_ROOT = os.path.join("data", "generated")

TITLE = Font(bold=True, size=13, color="1F3B57")
HDR = Font(bold=True, size=9, color="FFFFFF")
BOLD = Font(bold=True, size=9)
NORM = Font(size=9)
HFILL = PatternFill("solid", fgColor="2E6C8E")
SUBFILL = PatternFill("solid", fgColor="DCE7F0")
thin = Side(style="thin", color="AEB8C4")
BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)
CTR = Alignment(horizontal="center", vertical="center", wrap_text=True)
LFT = Alignment(horizontal="left", vertical="center", wrap_text=True)


def _grid(ws, headers, rows, start=1, widths=None):
    for j, h in enumerate(headers, 1):
        c = ws.cell(start, j, h); c.font = HDR; c.fill = HFILL; c.alignment = CTR; c.border = BORDER
    for i, row in enumerate(rows, start + 1):
        for j, v in enumerate(row, 1):
            c = ws.cell(i, j, v); c.font = NORM; c.alignment = LFT; c.border = BORDER
    if widths:
        for j, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(j)].width = w


def _titleblock(ws, title, doc_no, part, span=8):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=span)
    t = ws.cell(1, 1, f"{part['supplier'].upper()}"); t.font = Font(bold=True, size=11, color="2E6C8E")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=span)
    d = ws.cell(2, 1, title); d.font = TITLE; d.alignment = CTR
    ws.cell(3, 1, f"Part: {part['part_name']} ({part['part_no']})").font = BOLD
    ws.cell(3, span, f"Doc: {doc_no}").font = NORM
    ws.cell(4, 1, f"Customer: {part['customer']}").font = NORM
    ws.cell(4, span, f"Material: {part['material']}").font = NORM
    return 6  # next free row


def gen_pfd(part, path):
    wb = Workbook(); ws = wb.active; ws.title = "PFD"
    r = _titleblock(ws, "PROCESS FLOW DIAGRAM", "PFD/" + part["part_no"], part, span=5)
    rows = [[o["op_no"], o["name"], o["machine"],
             "; ".join(c["name"] for c in o.get("product_chars", [])),
             o.get("cp_ref", "")] for o in part["operations"]]
    _grid(ws, ["Op No", "Operation Description", "Machine / Device",
               "Product Characteristics", "Control Plan Ref"], rows, r, [8, 30, 22, 40, 16])
    wb.save(path)


def gen_pfmea(part, path):
    wb = Workbook(); ws = wb.active; ws.title = "PFMEA"
    r = _titleblock(ws, "PROCESS FMEA", "PFMEA/" + part["part_no"], part, span=11)
    opname = {o["op_no"]: o["name"] for o in part["operations"]}
    rows = [[f["op_no"], opname.get(f["op_no"], ""), f["failure_mode"], f["effect"],
             f["severity"], f["cause"], f["occurrence"], f["detection"], f["detection_val"],
             f["rpn"], f.get("recommended_action", "")] for f in part["failure_modes"]]
    _grid(ws, ["Op", "Function", "Potential Failure Mode", "Effect", "Sev", "Cause",
               "Occ", "Current Control (Detection)", "Det", "RPN", "Recommended Action"],
          rows, r, [6, 18, 22, 24, 5, 22, 5, 24, 5, 6, 24])
    wb.save(path)


def gen_control_plan(part, path):
    wb = Workbook(); ws = wb.active; ws.title = "Control Plan"
    r = _titleblock(ws, "CONTROL PLAN", part["operations"][0].get("cp_ref", "CP"), part, span=8)
    rows = []
    for o in part["operations"]:
        for ch in o.get("product_chars", []) + o.get("process_chars", []):
            rows.append([o["op_no"], o["name"], o["machine"], ch["name"],
                         ch.get("special", ""), ch.get("spec", ""),
                         ch.get("unit", ""), "Per lot"])
    _grid(ws, ["Op", "Process", "Machine/Gauge", "Characteristic", "Special",
               "Specification / Tolerance", "Unit", "Frequency"], rows, r, [6, 22, 18, 26, 8, 26, 8, 12])
    wb.save(path)


def gen_inward(part, rec, path):
    wb = Workbook(); ws = wb.active; ws.title = "Inward Inspection"
    r = _titleblock(ws, "INWARD INSPECTION REPORT", "SI-F-QA-05", part, span=8)
    ws.cell(r, 1, f"Supplier: {rec['supplier']}").font = NORM
    ws.cell(r, 5, f"Date: {rec['date']}").font = NORM
    r += 1
    rows = [[i + 1, p["param"], p["spec"], p["mode"], *p["obs"], p["status"]]
            for i, p in enumerate(rec["params"])]
    _grid(ws, ["Sr", "Parameter", "Specification", "Mode", "Obs1", "Obs2", "Obs3", "Status"],
          rows, r, [5, 26, 22, 14, 8, 8, 8, 12])
    wb.save(path)


def gen_inprocess(part, rec, path):
    wb = Workbook(); ws = wb.active; ws.title = "In-Process IR"
    r = _titleblock(ws, "IN-PROCESS INSPECTION REPORT", "SI-F-QA-04", part, span=9)
    ws.cell(r, 1, f"Operation {rec['op_no']} · {rec['machine']}").font = BOLD
    ws.cell(r, 6, f"Date: {rec['date']}").font = NORM
    r += 1
    rows = [[i + 1, p["param"], p["spec"], p["mode"], p["first"], *p["obs"], p["last"]]
            for i, p in enumerate(rec["params"])]
    _grid(ws, ["Sr", "Parameter", "Specification", "Mode", "First",
               "1", "2", "3", "4", "5", "Last"], rows, r, [5, 22, 18, 12, 8, 7, 7, 7, 7, 7, 8])
    wb.save(path)


def gen_pdir(part, rec, path):
    wb = Workbook(); ws = wb.active; ws.title = "PDIR"
    r = _titleblock(ws, "PRE-DESPATCH INSPECTION REPORT", "SI-F-QA-03", part, span=9)
    ws.cell(r, 1, f"Invoice: {rec['invoice']}").font = NORM
    ws.cell(r, 6, f"Qty: {rec['qty']}   Status: {rec['status']}").font = BOLD
    r += 1
    rows = [[i + 1, c["char"], c["spec"], c["mode"], *c["obs"]]
            for i, c in enumerate(rec["chars"])]
    _grid(ws, ["Sl", "Characteristic", "Specification", "Mode", "1", "2", "3", "4", "5"],
          rows, r, [5, 24, 20, 14, 8, 8, 8, 8, 8])
    wb.save(path)


def gen_ncr(part, ncr_rows, path):
    wb = Workbook(); ws = wb.active; ws.title = "INT NC & CAR"
    r = _titleblock(ws, "NON-CONFORMANCE & CORRECTIVE ACTION", "SI-F-QA-01", part, span=9)
    rows = [[n["sl"], n["part"], n["act_qty"], n["inprocess_rej"], n["pdi_rej"],
             f'{n["rej_pct"]}%', n["reason"], n["corrective_action"], n["status"]]
            for n in ncr_rows]
    _grid(ws, ["SL", "Part", "Act Qty", "In-proc Rej", "PDI Rej", "Rej %",
               "Reason", "Corrective Action", "Status"], rows, r, [5, 16, 10, 11, 9, 8, 30, 34, 10])
    wb.save(path)


def gen_capa(part, capa, path):
    wb = Workbook(); ws = wb.active; ws.title = "8D CAPA"
    ws.merge_cells("A1:H1"); c = ws.cell(1, 1, "8D REPORT — CAPA ANALYSIS"); c.font = TITLE; c.alignment = CTR
    ws.column_dimensions["A"].width = 26
    for j in range(2, 9):
        ws.column_dimensions[get_column_letter(j)].width = 16

    def kv(row, label, value):
        a = ws.cell(row, 1, label); a.font = BOLD; a.fill = SUBFILL; a.border = BORDER; a.alignment = LFT
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=8)
        b = ws.cell(row, 2, value); b.font = NORM; b.alignment = LFT; b.border = BORDER
    rr = 3
    kv(rr, "Notification No.", capa["notification_no"]); rr += 1
    kv(rr, "Date", capa["date"]); rr += 1
    kv(rr, "Part / Customer", f'{capa["part_name"]} ({part["part_no"]}) · {capa["customer"]}'); rr += 1
    kv(rr, "D1 Team", ", ".join(capa["team"])); rr += 1
    p = capa["problem"]
    kv(rr, "D2 What", p["what"]); rr += 1
    kv(rr, "D2 Why", p["why"]); rr += 1
    kv(rr, "D2 When / Where", f'{p["when"]} · {p["where"]}'); rr += 1
    kv(rr, "D2 How / How many", f'{p["how"]} · {p["how_many"]}'); rr += 1
    kv(rr, "D3 Containment", capa["containment"]); rr += 1
    kv(rr, "D4 Root Cause", capa["root_cause"]); rr += 1
    kv(rr, "D5 Corrective Action", capa["corrective_action"]); rr += 1
    kv(rr, "D7 Preventive Action", capa["preventive_action"]); rr += 1
    kv(rr, "D8 Status", capa["status"]); rr += 1
    wb.save(path)


def gen_wi_docx(part, path):
    """English work instruction (docx) for parts that have one authored in English."""
    from docx import Document
    wi = part.get("wi")
    if not wi or wi.get("language") != "English":
        return False
    doc = Document()
    doc.add_heading(f"{part['supplier']} — Work Instruction", level=1)
    doc.add_heading(wi["title"], level=2)
    doc.add_paragraph(f"Part: {part['part_name']} ({part['part_no']})   Operation: {wi['op_no']}")
    for i, step in enumerate(wi["steps_en"], 1):
        doc.add_paragraph(f"{i}. {step}")
    doc.save(path)
    return True


def generate_all():
    made = []
    for pno, part in PARTS.items():
        d = os.path.join(OUT_ROOT, pno.replace("/", "-"))
        os.makedirs(d, exist_ok=True)
        gen_pfd(part, os.path.join(d, f"{pno}_PFD.xlsx"))
        gen_pfmea(part, os.path.join(d, f"{pno}_PFMEA.xlsx"))
        gen_control_plan(part, os.path.join(d, f"{pno}_ControlPlan.xlsx"))
        if pno in INWARD:
            gen_inward(part, INWARD[pno], os.path.join(d, f"{pno}_InwardInspection.xlsx"))
        if pno in INPROCESS:
            gen_inprocess(part, INPROCESS[pno], os.path.join(d, f"{pno}_InProcessIR.xlsx"))
        if pno in PDIR:
            gen_pdir(part, PDIR[pno], os.path.join(d, f"{pno}_PDIR.xlsx"))
        ncr_rows = [n for n in NCRS if n["part_no"] == pno]
        if ncr_rows:
            gen_ncr(part, ncr_rows, os.path.join(d, f"{pno}_NCR_CAR.xlsx"))
        for capa in [c for c in CAPAS if c["part_no"] == pno]:
            gen_capa(part, capa, os.path.join(d, f"{pno}_8D_{capa['notification_no']}.xlsx"))
        if gen_wi_docx(part, os.path.join(d, f"{pno}_WorkInstruction.docx")):
            pass
        made.append((pno, len(os.listdir(d))))
    return made


if __name__ == "__main__":
    for pno, n in generate_all():
        print(f"{pno}: {n} files")
